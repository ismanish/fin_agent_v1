# aqrr_word_generate.py
import os
import io
import json
import re
import requests
import numpy as np
import pandas as pd
from datetime import datetime
import argparse
import sys

# Word document generation imports
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from fastapi import APIRouter, FastAPI, HTTPException, Body
from fastapi.responses import StreamingResponse
from io import BytesIO

router = APIRouter()
app = FastAPI(title="Word API")
app.include_router(router, prefix="/word")


@router.get('/get_companies')
def get_companies():
    """Dynamically lists company folders."""
    base_folder = os.getenv('COMPANY_DATA_FOLDER', 'company_data')
    try:
        companies = [
            d for d in os.listdir(base_folder)
            if os.path.isdir(os.path.join(base_folder, d))
        ]
        return companies
    except FileNotFoundError:
        raise HTTPException(status_code=404, detail="Company data folder not found.")


@router.post('/aqrr_word')
def generate_word(data: dict = Body(...)):
    """Generates a Word document from company data."""
    company_name = data.get('company')
    if not company_name:
        raise HTTPException(status_code=400, detail="No company name provided.")

    base_folder = os.getenv('COMPANY_DATA_FOLDER', 'company_data')
    company_path = os.path.join(base_folder, company_name)
    if not os.path.exists(company_path):
        raise HTTPException(status_code=404, detail="Company folder not found.")

    # Find the data file (csv, excel, or json)
    data_file = None
    for filename in os.listdir(company_path):
        if filename.endswith(('.csv', '.xlsx', '.json')):
            data_file = os.path.join(company_path, filename)
            break

    if not data_file:
        raise HTTPException(status_code=404, detail="Data file (csv, xlsx, or json) not found.")

    # Load data from the file
    try:
        if data_file.endswith('.csv'):
            df = pd.read_csv(data_file)
            # Replace NaN values with empty strings for CSV files
            df = df.replace({np.nan: ''})
        elif data_file.endswith('.xlsx'):
            # Explicitly specify the sheet name
            df = pd.read_excel(data_file, sheet_name='Essence Table')
        elif data_file.endswith('.json'):
            with open(data_file, 'r') as f:
                json_data = json.load(f)
            df = json_to_dataframe(json_data)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error reading data file or worksheet: {e}")

    # Handle case where the data file or sheet is empty
    if df.empty:
        raise HTTPException(status_code=500, detail="Data file or specified worksheet is empty. Cannot create a table.")

    # Load statement analysis text
    analysis_file = os.path.join(company_path, 'statement_analysis.txt')
    if os.path.exists(analysis_file):
        with open(analysis_file, 'r') as f:
            analysis_text = f.read()
    else:
        analysis_text = "No statement analysis text found."

    # Generate the Word document
    buffer = BytesIO()
    
    # Create Word document
    doc = create_word_document(df, analysis_text, data_file, company_name)
    
    # Save the document to the buffer
    doc.save(buffer)
    buffer.seek(0)

    headers = {"Content-Disposition": f"attachment; filename={company_name}_report.docx"}
    return StreamingResponse(buffer, media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document', headers=headers)

# Helper functions from PDF generator
def get_company_title_from_ticker(ticker: str, mapping_path: str = os.path.join('static', 'company_ticker.json')) -> str:
    """Return company title/name for a given ticker using static/company_ticker.json.
    Falls back to ticker if not found or file missing.
    """
    try:
        with open(mapping_path, 'r') as f:
            mapping = json.load(f)
        t_upper = ticker.upper()
        for _, entry in mapping.items():
            if isinstance(entry, dict) and entry.get('ticker', '').upper() == t_upper:
                return entry.get('title') or t_upper
    except Exception:
        pass


def set_table_indent(table, inches: float = 0.0):
    """Set table left indent explicitly to avoid unexpected horizontal offset."""
    try:
        tbl = table._tbl
        tblPr = tbl.tblPr
        if tblPr is None:
            tblPr = OxmlElement('w:tblPr')
            tbl.append(tblPr)
        tblInd = tblPr.find(qn('w:tblInd'))
        if tblInd is None:
            tblInd = OxmlElement('w:tblInd')
            tblPr.append(tblInd)
        twips = int(inches * 1440)
        tblInd.set(qn('w:w'), str(twips))
        tblInd.set(qn('w:type'), 'dxa')
    except Exception:
        pass


def set_column_preferred_width(table, col_idx: int, width_in: float):
    """Force a column's preferred width for all cells (sets tcW)."""
    try:
        width_twips = int(width_in * 1440)
        column = table.columns[col_idx]
        column.width = Inches(width_in)
        for cell in column.cells:
            cell.width = Inches(width_in)
            tcPr = cell._element.tcPr
            if tcPr is None:
                tcPr = OxmlElement('w:tcPr')
                cell._element.append(tcPr)
            tcW = tcPr.find(qn('w:tcW'))
            if tcW is None:
                tcW = OxmlElement('w:tcW')
                tcPr.append(tcW)
            tcW.set(qn('w:w'), str(width_twips))
            tcW.set(qn('w:type'), 'dxa')
    except Exception:
        pass


def get_company_title_from_sec(ticker: str,
                               url: str = 'https://www.sec.gov/files/company_tickers.json') -> str | None:
    """Attempt to resolve company title from SEC's public company_tickers.json.
    Returns the company title or None if not found/failed.
    """
    try:
        resp = requests.get(url, timeout=30)
        if resp.status_code != 200:
            return None
        data = resp.json()
        t_up = str(ticker).upper().strip()
        if isinstance(data, dict):
            for _, v in data.items():
                try:
                    if isinstance(v, dict) and str(v.get('ticker', '')).upper().strip() == t_up:
                        title = v.get('title')
                        return str(title) if title else None
                except Exception:
                    continue
        return None
    except Exception:
        return None


def current_quarter_index(reference: datetime | None = None) -> int:
    ref = reference or datetime.now()
    m = ref.month
    if m <= 3:
        return 1
    if m <= 6:
        return 2
    if m <= 9:
        return 3
    return 4


def quarter_end_label_for_year(year: int, reference: datetime | None = None) -> str:
    """Return the quarter-end label like '3/31/25' based on the current quarter for the given year."""
    q = current_quarter_index(reference)
    md = [(3, 31), (6, 30), (9, 30), (12, 31)][q - 1]
    mm, dd = md
    yy = year % 100
    # return f"{mm}/{dd}/{yy:02d}"
    return f"03/31/{yy:02d}"


def format_number_for_display(val):
    """Format numbers for display in HFA table: remove 000s and format negatives with parentheses"""
    try:
        if val is None or val == "" or (isinstance(val, str) and val.strip() == "-"):
            return "-"
        
        # Handle percentage values differently
        if isinstance(val, str) and '%' in val:
            return val
            
        # Convert to float and divide by 1000 to remove 000s
        f = float(val) / 1000
        
        # Format negative numbers with parentheses
        if f < 0:
            # Remove the negative sign and wrap in parentheses
            if abs(f - int(f)) < 1e-6:  # Integer
                return f"({int(abs(f)):,})"
            else:  # Float with decimals
                return f"({abs(f):,.1f})"
        else:
            # Positive numbers
            if abs(f - int(f)) < 1e-6:  # Integer
                return f"{int(f):,}"
            else:  # Float with decimals
                return f"{f:,.1f}"
    except Exception:
        # If conversion fails, return as is
        return str(val)


def format_ratio_to_two_decimals(val):
    """Format ratio strings like '3.3x' to two decimals: '3.30x'. Leaves non-ratio values unchanged."""
    try:
        if val is None:
            return ""
        s = str(val).strip()
        if s in ("", "-", "–", "—"):
            return s
        # Detect trailing 'x' ratio suffix
        has_x = s.lower().endswith('x')
        core = s[:-1].strip() if has_x else s
        # Remove commas
        core = core.replace(',', '')
        num = float(core)
        formatted = f"{num:.2f}"
        return f"{formatted}x" if has_x else formatted
    except Exception:
        return str(val)


def flatten_json(nested_json, prefix='', separator='_'):
    """
    Flatten a nested JSON structure into a flat dictionary.
    Improved to handle complex nested structures including lists of objects.
    """
    flattened = {}
    for key, value in nested_json.items():
        if isinstance(value, dict):
            flattened.update(flatten_json(value, f"{prefix}{key}{separator}", separator))
        elif isinstance(value, list):
            if value and all(isinstance(item, dict) for item in value):
                for i, item in enumerate(value):
                    flattened.update(flatten_json(item, f"{prefix}{key}{separator}{i}{separator}", separator))
            else:
                flattened[f"{prefix}{key}"] = json.dumps(value)
        else:
            flattened[f"{prefix}{key}"] = value
    return flattened


def json_to_dataframe(json_data):
    """
    Convert JSON data to a pandas DataFrame.
    Handles both array of objects and single object formats.
    Improved to better handle complex nested structures.
    """
    if isinstance(json_data, list):
        if not json_data:
            return pd.DataFrame()

        if all(isinstance(item, dict) for item in json_data):
            try:
                return pd.json_normalize(json_data)
            except Exception:
                flattened_data = [flatten_json(item) for item in json_data]
                return pd.DataFrame(flattened_data)
        else:
            return pd.DataFrame(json_data, columns=['Value'])
    elif isinstance(json_data, dict):
        try:
            return pd.json_normalize([json_data])
        except Exception:
            flattened_data = flatten_json(json_data)
            return pd.DataFrame([flattened_data])
    else:
        return pd.DataFrame([json_data], columns=['Value'])


def set_cell_border(cell, **kwargs):
    """
    Set cell border properties in a Word table.
    """
    tc = cell._element.tcPr
    if tc is None:
        tc = OxmlElement('w:tcPr')
        cell._element.append(tc)

    # Set border properties
    for key, value in kwargs.items():
        tag = 'w:{}'.format(key)
        element = tc.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            tc.append(element)

        # Set val attribute
        element.set(qn('w:val'), value)


def set_cell_background(cell, color, text_color=None):
    """
    Set cell background color and optionally text color.
    """
    tc = cell._element.tcPr
    if tc is None:
        tc = OxmlElement('w:tcPr')
        cell._element.append(tc)
    
    # Set background color
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    tc.append(shading)
    
    # Set text color if provided
    if text_color and len(cell.paragraphs) > 0 and len(cell.paragraphs[0].runs) > 0:
        run = cell.paragraphs[0].runs[0]
        run.font.color.rgb = text_color


def set_table_fixed_width(table, width_in: float):
    """Force a table to a fixed width by setting tblW and a fixed layout.
    This prevents Word from shrinking the table when cells are mostly empty.
    """
    try:
        tbl = table._tbl
        tblPr = tbl.tblPr
        if tblPr is None:
            tblPr = OxmlElement('w:tblPr')
            tbl.append(tblPr)
        # Set fixed layout
        tblLayout = tblPr.find(qn('w:tblLayout'))
        if tblLayout is None:
            tblLayout = OxmlElement('w:tblLayout')
            tblPr.append(tblLayout)
        tblLayout.set(qn('w:type'), 'fixed')
        # Set width in twips (1 inch = 1440 twips)
        twips = int(width_in * 1440)
        tblW = tblPr.find(qn('w:tblW'))
        if tblW is None:
            tblW = OxmlElement('w:tblW')
            tblPr.append(tblW)
        tblW.set(qn('w:w'), str(twips))
        tblW.set(qn('w:type'), 'dxa')
    except Exception:
        pass


def add_header_footer(doc, company_name=None):
    """
    Add header and footer to the Word document.
    """
    # Add header
    section = doc.sections[0]
    header = section.header
    
    # Use the existing paragraph for the header
    header_para = header.paragraphs[0]
    header_para.text = ""
    
    # Add three runs with different alignments
    left_run = header_para.add_run("PGIM Private Capital")
    left_run.font.name = 'Calibri'
    left_run.font.size = Pt(12)
    left_run.font.bold = True
    
    # Add tab and center text
    header_para.add_run("\t")
    center_run = header_para.add_run("Annual Quality Rating Review")
    center_run.font.name = 'Calibri'
    center_run.font.size = Pt(12)
    center_run.font.bold = True
    
    # Add tab and right-aligned date
    header_para.add_run("\t")
    # right_run = header_para.add_run(datetime.now().strftime("%m/%d/%Y"))
    # right_run.font.name = 'Calibri'
    # right_run.font.size = Pt(12)
    # right_run.font.bold = True
    
    # Set tab stops for proper alignment
    tab_stops = header_para.paragraph_format.tab_stops
    tab_stops.add_tab_stop(Inches(4.135), WD_ALIGN_PARAGRAPH.CENTER)  # Center of page
    tab_stops.add_tab_stop(Inches(8.27), WD_ALIGN_PARAGRAPH.RIGHT)   # Right side
    
    # Add footer with page numbers
    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.text = "Page "
    footer_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    # Add page number field
    run = footer_para.add_run()
    fld_char = OxmlElement('w:fldChar')
    fld_char.set(qn('w:fldCharType'), 'begin')
    run._element.append(fld_char)
    
    instr_text = OxmlElement('w:instrText')
    instr_text.set(qn('xml:space'), 'preserve')
    instr_text.text = "PAGE"
    run._element.append(instr_text)
    
    fld_char = OxmlElement('w:fldChar')
    fld_char.set(qn('w:fldCharType'), 'end')
    run._element.append(fld_char)
    
    # Format footer text
    run = footer_para.runs[0]
    run.font.name = 'Calibri'
    run.font.size = Pt(9)


def create_word_document(df, analysis_text, data_file, company_name=None):
    """
    Create a Word document with the same formatting as the PDF.
    """
    doc = Document()
    
    # Set document properties
    section = doc.sections[0]
    section.page_width = Inches(8.27)  # A4 width
    section.page_height = Inches(11.69)  # A4 height
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    
    # Add header and footer
    add_header_footer(doc, company_name)
    
    # Create table from DataFrame
    if data_file.endswith('.csv'):
        # Special handling for CSV files to create a two-row header
        if len(df) >= 2:
            # Extract the first two rows for headers
            header_row1 = df.columns.tolist()
            header_row2 = df.iloc[0].tolist()
            
            # Remove the first row from the dataframe as it's now part of the header
            df = df.iloc[1:].reset_index(drop=True)
            
            # Create table with appropriate dimensions
            num_rows = len(df) + 2  # +2 for the two header rows
            num_cols = len(df.columns)
            table = doc.add_table(rows=num_rows, cols=num_cols)
            table.style = 'Table Grid'
            
            # Fill in the header rows
            for i, cell_text in enumerate(header_row1):
                cell = table.cell(0, i)
                cell.text = str(cell_text)
                # Format header cell
                cell_para = cell.paragraphs[0]
                cell_para.alignment = WD_ALIGN_PARAGRAPH.CENTER if i > 0 else WD_ALIGN_PARAGRAPH.LEFT
                run = cell_para.runs[0]
                run.font.name = 'Calibri'
                run.font.size = Pt(8)
                run.font.bold = True
                # Set background color
                shading = OxmlElement('w:shd')
                shading.set(qn('w:fill'), "D3D3D3")  # Light gray
                cell._element.tcPr.append(shading)
            
            for i, cell_text in enumerate(header_row2):
                cell = table.cell(1, i)
                cell.text = str(cell_text) if cell_text != '' else ''
                # Format header cell
                cell_para = cell.paragraphs[0]
                cell_para.alignment = WD_ALIGN_PARAGRAPH.CENTER if i > 0 else WD_ALIGN_PARAGRAPH.LEFT
                run = cell_para.runs[0]
                run.font.name = 'Calibri'
                run.font.size = Pt(8)
                run.font.bold = True
                # Set background color
                shading = OxmlElement('w:shd')
                shading.set(qn('w:fill'), "D3D3D3")  # Light gray
                cell._element.tcPr.append(shading)
            
            # Fill in the data rows
            # Track the Word row index of the 'Key Financial Ratios:' header (two header rows offset)
            ratio_header_word_row_idx = None
            for i, row in enumerate(df.values):
                for j, cell_text in enumerate(row):
                    cell = table.cell(i + 2, j)  # +2 to account for the header rows
                    cell.text = str(cell_text) if cell_text != '' else ''
                    # Format data cell
                    cell_para = cell.paragraphs[0]
                    cell_para.alignment = WD_ALIGN_PARAGRAPH.CENTER if j > 0 else WD_ALIGN_PARAGRAPH.LEFT
                    run = cell_para.runs[0]
                    run.font.name = 'Calibri'
                    run.font.size = Pt(8)  # Increased font size
                    
                    # Check for special formatting
                    first_cell_value = str(df.iloc[i, 0]) if len(df.columns) > 0 else ""
                    exact_keywords = ["Total Debt", "Total Debt + COLs", "Book Capitalization", "Market Capitalization"]
                    
                    if first_cell_value in exact_keywords:
                        # Make the entire row bold
                        run.font.bold = True
                        # Add indent to first column if this is the first cell
                        if j == 0:
                            cell.text = f"   {first_cell_value}"
                    
                    # Special handling for "Key Financial Ratios:"
                    elif first_cell_value == "Key Financial Ratios:":
                        if j == 0:
                            # Span the cell across all columns
                            for k in range(1, num_cols):
                                table.cell(i + 2, 0).merge(table.cell(i + 2, k))
                            # Format the cell
                            cell.text = first_cell_value
                            cell_para = cell.paragraphs[0]
                            cell_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            run = cell_para.runs[0]
                            run.font.name = 'Calibri'
                            run.font.size = Pt(8)
                            run.font.bold = True
                            run.italic = True
                            # Set background color to light gray
                            shading = OxmlElement('w:shd')
                            shading.set(qn('w:fill'), "D3D3D3")  # Light gray
                            cell._element.tcPr.append(shading)
            
            # Set column widths
            # First column gets 45% of the table width, remaining columns share the rest
            table.autofit = False
            table.allow_autofit = False
            for i, column in enumerate(table.columns):
                if i == 0:
                    column.width = Inches(3.5)  # 45% of ~7.27 inches (A4 width minus margins)
                else:
                    column.width = Inches(3.77 / (num_cols - 1))  # Remaining width divided among other columns
        
        else:
            # If there aren't enough rows for a two-row header, use a single row header
            num_rows = len(df) + 1  # +1 for the header row
            num_cols = len(df.columns)
            table = doc.add_table(rows=num_rows, cols=num_cols)
            table.style = 'Table Grid'
            
            # Fill in the header row
            for i, cell_text in enumerate(df.columns):
                cell = table.cell(0, i)
                cell.text = str(cell_text)
                # Format header cell
                cell_para = cell.paragraphs[0]
                cell_para.alignment = WD_ALIGN_PARAGRAPH.CENTER if i > 0 else WD_ALIGN_PARAGRAPH.LEFT
                run = cell_para.runs[0]
                run.font.name = 'Calibri'
                run.font.size = Pt(8)
                run.font.bold = True
                # Set background color
                shading = OxmlElement('w:shd')
                shading.set(qn('w:fill'), "D3D3D3")  # Light gray
                cell._element.tcPr.append(shading)
            
            # Fill in the data rows
            for i, row in enumerate(df.values):
                for j, cell_text in enumerate(row):
                    cell = table.cell(i + 1, j)  # +1 to account for the header row
                    cell.text = str(cell_text) if cell_text != '' else ''
                    # Format data cell
                    cell_para = cell.paragraphs[0]
                    cell_para.alignment = WD_ALIGN_PARAGRAPH.CENTER if j > 0 else WD_ALIGN_PARAGRAPH.LEFT
                    run = cell_para.runs[0]
                    run.font.name = 'Calibri'
                    run.font.size = Pt(8)  # Increased font size
            
            # Set column widths
            table.autofit = False
            table.allow_autofit = False
            for i, column in enumerate(table.columns):
                if i == 0:
                    column.width = Inches(2.5)  # 30% of ~8.27 inches (A4 width)
                else:
                    column.width = Inches(4.77 / (num_cols - 1))  # Remaining width divided among other columns
    
    else:
        # Original handling for non-CSV files
        num_rows = len(df) + 1  # +1 for the header row
        num_cols = len(df.columns)
        table = doc.add_table(rows=num_rows, cols=num_cols)
        table.style = 'Table Grid'
        
        # Fill in the header row
        for i, cell_text in enumerate(df.columns):
            cell = table.cell(0, i)
            cell.text = str(cell_text)
            # Format header cell
            cell_para = cell.paragraphs[0]
            cell_para.alignment = WD_ALIGN_PARAGRAPH.CENTER if i > 0 else WD_ALIGN_PARAGRAPH.LEFT
            run = cell_para.runs[0]
            run.font.name = 'Calibri'
            run.font.size = Pt(8)
            run.font.bold = True
            # Set background color
            shading = OxmlElement('w:shd')
            shading.set(qn('w:fill'), "D3D3D3")  # Light gray
            cell._element.tcPr.append(shading)
        
        # Fill in the data rows
        for i, row in enumerate(df.values):
            for j, cell_text in enumerate(row):
                cell = table.cell(i + 1, j)  # +1 to account for the header row
                cell.text = str(cell_text) if cell_text != '' else ''
                # Format data cell
                cell_para = cell.paragraphs[0]
                cell_para.alignment = WD_ALIGN_PARAGRAPH.CENTER if j > 0 else WD_ALIGN_PARAGRAPH.LEFT
                run = cell_para.runs[0]
                run.font.name = 'Calibri'
                run.font.size = Pt(8)  # Increased font size
                
                # Styling based on first column content
                first_cell_value = str(df.iloc[i, 0]) if len(df.columns) > 0 else ""
                if any(keyword in first_cell_value for keyword in ["Total Debt", "Book Capitalization", "Market Capitalization"]):
                    # Make the entire row bold
                    run.font.bold = True
                    # Add indent to first column if this is the first cell
                    if j == 0:
                        cell.text = f"   {first_cell_value}"
                        run = cell.paragraphs[0].runs[0]
                        run.font.name = 'Calibri'
                        run.font.size = Pt(8)  # Increased font size
                        run.font.bold = True
        
        # Set column widths
        table.autofit = False
        table.allow_autofit = False
        for i, column in enumerate(table.columns):
            if i == 0:
                column.width = Inches(1.65)  # 20% of ~8.27 inches (A4 width)
            else:
                column.width = Inches(5.62 / (num_cols - 1))  # Remaining width divided among other columns
    
    # Add a page break after the table
    doc.add_paragraph().add_run().add_break()
    
    # Add text from statement analysis
    doc.add_heading("Statement Analysis", level=1)
    for line in analysis_text.split('\n'):
        doc.add_paragraph(line)
    
    return doc


def build_word_bytes_from_ticker(ticker: str,
                                hfa_dir: str = os.path.join('output', 'json', 'hfa_output'),
                                fsa_dir: str = os.path.join('output', 'json', 'financial_analysis'),
                                prefetched_data: dict = None) -> bytes:
    """
    Build the Word document for a given ticker by calling the HFA API and using its rows:
    - HFA table data from: POST {BASE_URL}/api/v1/hfa with body {"ticker": TICKER}
      BASE_URL is taken from env APP_BASE_URL (default http://127.0.0.1:9259)
    - Financial Statement Analysis from: output/json/financial_analysis/{TICKER}_FSA.json
    - CAP table data from: POST {BASE_URL}/api/v1/cap-table with body {"ticker": TICKER}
    - COMP table data from: POST {BASE_URL}/api/v1/comp with body {"ticker": TICKER}
    Returns raw Word document bytes.
    """
    if not ticker:
        raise ValueError("No ticker provided.")

    if prefetched_data:
        # Use pre-fetched data
        hfa_rows = prefetched_data['hfa_rows']
        cap_json = prefetched_data['cap_json']
        comp_rows = prefetched_data['comp_rows']
        fsa_data = prefetched_data['fsa_data']
        credit_data = prefetched_data.get('credit_data')
        company_exposure = prefetched_data.get('company_exposure')
    else:
        # Resolve FSA input JSON path
        fsa_path = os.path.join(fsa_dir, f"{ticker}_FSA.json")
        # Call HFA API to get rows for the table
        api_base = os.getenv('APP_BASE_URL', 'http://127.0.0.1:9259')
        api_url = f"{api_base.rstrip('/')}/api/v1/hfa"
        try:
            resp = requests.post(api_url, json={"ticker": ticker}, timeout=300)
        except requests.RequestException as e:
            raise RuntimeError(f"Failed to call HFA API at {api_url}: {e}")
        if resp.status_code != 200:
            try:
                err_detail = resp.json()
            except Exception:
                err_detail = resp.text
            raise RuntimeError(f"HFA API returned {resp.status_code}: {err_detail}")
        try:
            payload = resp.json()
        except Exception as e:
            raise RuntimeError(f"Invalid JSON from HFA API: {e}")
        hfa_rows = payload.get("rows")
        if not isinstance(hfa_rows, list) or not hfa_rows:
            raise RuntimeError("HFA API response missing 'rows' list with data")

        # Fetch Credit Risk Metrics data (non-fatal)
        credit_data = None
        try:
            credit_url = f"{api_base.rstrip('/')}/api/v1/credit_table"
            credit_resp = requests.post(credit_url, json={"ticker": ticker}, timeout=300)
            if credit_resp.status_code == 200:
                try:
                    credit_payload = credit_resp.json()
                    if isinstance(credit_payload, dict):
                        credit_data = credit_payload.get("json_data")
                        # Fallback: parse raw JSON string if provided by API
                        if credit_data is None and isinstance(credit_payload.get("json_data_raw"), str):
                            raw = credit_payload.get("json_data_raw")
                            def _try_parse_json_text(s: str):
                                try:
                                    return json.loads(s)
                                except Exception:
                                    # sanitize and retry: remove trailing commas and trim to outer braces
                                    s2 = s.strip()
                                    if s2.startswith("```"):
                                        s2 = s2.strip('`')
                                    s2 = re.sub(r",\s*([}\\]])", r"\1", s2)
                                    if '{' in s2 and '}' in s2:
                                        s2 = s2[s2.find('{'): s2.rfind('}') + 1]
                                    try:
                                        return json.loads(s2)
                                    except Exception:
                                        return None
                            credit_data = _try_parse_json_text(raw)
                except Exception:
                    credit_data = None
        except Exception:
            credit_data = None

        # Fetch Company Exposure Details table (non-fatal)
        company_exposure = None
        try:
            company_url = f"{api_base.rstrip('/')}/api/v1/company-table"
            company_resp = requests.post(company_url, json={"ticker": ticker}, timeout=120)
            if company_resp.status_code == 200:
                try:
                    company_payload = company_resp.json()
                    if isinstance(company_payload, dict):
                        company_exposure = company_payload.get("table")
                except Exception:
                    company_exposure = None
        except Exception:
            company_exposure = None

        # Fetch CAP table JSON from API (non-fatal if unavailable)
        cap_json = None
        try:
            cap_url = f"{api_base.rstrip('/')}/api/v1/cap-table"
            cap_resp = requests.post(cap_url, json={"ticker": ticker}, timeout=300)
            if cap_resp.status_code == 200:
                try:
                    cap_payload = cap_resp.json()
                    if isinstance(cap_payload, dict):
                        cap_json = cap_payload.get("json_data")
                        # Fallback: parse raw JSON string if provided by API
                        if cap_json is None and isinstance(cap_payload.get("json_data_raw"), str):
                            raw = cap_payload.get("json_data_raw")
                            def _try_parse_json_text(s: str):
                                try:
                                    return json.loads(s)
                                except Exception:
                                    # sanitize and retry: remove trailing commas and trim to outer braces
                                    s2 = s.strip()
                                    if s2.startswith("```"):
                                        s2 = s2.strip('`')
                                    s2 = re.sub(r",\s*([}\]])", r"\1", s2)
                                    if '{' in s2 and '}' in s2:
                                        s2 = s2[s2.find('{'): s2.rfind('}') + 1]
                                    try:
                                        return json.loads(s2)
                                    except Exception:
                                        return None
                            cap_json = _try_parse_json_text(raw)
                except Exception:
                    cap_json = None
        except Exception:
            cap_json = None
        # Fetch COMP rows from API (non-fatal if unavailable)
        comp_rows = None
        try:
            comp_url = f"{api_base.rstrip('/')}/api/v1/comp"
            comp_resp = requests.post(comp_url, json={"ticker": ticker}, timeout=300)
            if comp_resp.status_code == 200:
                try:
                    comp_payload = comp_resp.json()
                    if isinstance(comp_payload, dict):
                        comp_rows = comp_payload.get("rows")
                except Exception:
                    comp_rows = None
        except Exception:
            comp_rows = None
        # Load FSA data if available
        fsa_data = None
        if os.path.exists(fsa_path):
            with open(fsa_path, 'r') as f:
                try:
                    fsa_data = json.load(f)
                except Exception:
                    fsa_data = None

    # Convert HFA rows to DataFrame
    df = json_to_dataframe(hfa_rows)
    # Clean NaNs for rendering
    df = df.replace({np.nan: '-'})
    df = df.replace({None: '-'})
    df = df.replace({0: '-'})
    # Format numbers per metric type.
    # IMPORTANT: Do NOT divide by 1000 for percentage or ratio rows; these are formatted later.
    percentage_metrics = {'% YoY Growth', '% Margin'}
    ratio_keywords = [
        'EBITDA / Int',
        'EBITDA / Interest',
        'EBITDAR / Interest',
        'EBITDAR / Interest + Rent',
        'Total Debt / EBITDA',
        'Total Debt + Leases / EBITDA',
        'Total Debt / Book',
        'Total Debt + Leases / Book',
    ]
    # Specific ratio metrics that should be displayed as percentages (not with 'x')
    percentage_ratio_metrics = {
        'Total Debt / Book Capital',
        'Total Debt + Leases / Book Capital',
    }
    if 'Metric' in df.columns:
        for i, row in df.iterrows():
            metric_name = str(row['Metric'])
            for col in df.columns:
                if col == 'Metric':
                    continue
                val = row[col]
                if val == '':
                    df.at[i, col] = ''
                    continue
                if (metric_name in percentage_metrics) or any(k in metric_name for k in ratio_keywords):
                    # keep raw, will format later
                    df.at[i, col] = val
                else:
                    df.at[i, col] = format_number_for_display(val)
    else:
        # Fallback if no Metric column exists
        for col in df.columns:
            df[col] = df[col].apply(lambda x: format_number_for_display(x) if x != '' else '')
    # Special formatting for percentage rows
    for i, row in enumerate(df.values):
        metric = str(row[0]) if row[0] != '' else ''
        if metric in ['% YoY Growth', '% Margin'] or metric in percentage_ratio_metrics:
            for j, val in enumerate(row):
                if j > 0 and val not in ['', '-']:
                    try:
                        # Format as percentage with one decimal place
                        num_val = float(val.replace('(', '-').replace(')', '').replace(',', '')) if isinstance(val, str) else float(val)
                        if num_val < 0:
                            df.iloc[i, j] = f'({abs(num_val):.1f}%)'
                        else:
                            df.iloc[i, j] = f'{num_val:.1f}%'
                    except:
                        # Keep as is if conversion fails
                        pass
    # Ensure 'Metric' is the first column if present
    if 'Metric' in df.columns:
        cols = df.columns.tolist()
        cols.remove('Metric')
        df = df[['Metric'] + cols]
    # Reorder columns into: Metric | years (asc) | YTD years (asc) | LTM years (asc)
    # Also capture groups to construct a two-row header later
    all_cols = df.columns.tolist()
    year_cols = [c for c in all_cols if isinstance(c, str) and c.isdigit() and len(c) == 4]
    try:
        year_cols_sorted = sorted(year_cols, key=lambda x: int(x))
    except Exception:
        year_cols_sorted = year_cols
    ytd_cols = [c for c in all_cols if isinstance(c, str) and c.startswith('YTD ')]
    try:
        ytd_cols_sorted = sorted(ytd_cols, key=lambda x: int(x.split()[1]))
    except Exception:
        ytd_cols_sorted = ytd_cols
    ltm_cols = [c for c in all_cols if isinstance(c, str) and c.startswith('LTM ')]
    try:
        ltm_cols_sorted = sorted(ltm_cols, key=lambda x: int(x.split()[1]))
    except Exception:
        ltm_cols_sorted = ltm_cols
    ordered_cols = ['Metric']
    ordered_cols += [c for c in year_cols_sorted if c in all_cols]
    ordered_cols += [c for c in ytd_cols_sorted if c in all_cols]
    ordered_cols += [c for c in ltm_cols_sorted if c in all_cols]
    # Only reorder if all expected columns present
    if all(c in all_cols for c in ordered_cols) and len(ordered_cols) == len(all_cols):
        df = df[ordered_cols]

    # Get company title
    company_title = get_company_title_from_ticker(ticker)

    # Generate the Word document in-memory
    buffer = BytesIO()
    
    # Create Word document
    doc = Document()
    
    # Set document properties
    section = doc.sections[0]
    section.page_width = Inches(8.27)  # A4 width
    section.page_height = Inches(11.69)  # A4 height
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    
    # Add header and footer
    add_header_footer(doc, company_title)
    
    # Company Exposure Details table (placed above Credit Merits/Risks)
    try:
        if isinstance(company_exposure, dict) and company_exposure:
            # Exact row-wise order to match screenshot (3 pairs per row)
            labels_grid = [
                ["Credit Exposure Name", "Public / Private", "Servicing Category"],
                ["iRisk Parent Name", "Headquarters", "PruScore"],
                ["Industry", "Region", "NAIC Designation"],
                ["Real Assets Category", "MD", "S&P / M / Fitch"],
                ["Stat. Country", "Team Leader", "Other Ratings"],
                ["Economic Risk Country", "Secondary", "Unqualified Audit"],
                ["Valuation Country", "Analyst", None],
            ]

            # Create 6-column table: (Label|Value) x 3
            rows_needed = len(labels_grid)
            doc.add_paragraph()  # spacing before
            det_table = doc.add_table(rows=rows_needed, cols=6)
            det_table.style = 'Table Grid'
            det_table.autofit = False
            det_table.allow_autofit = False

            # Column widths to fit within ~7.27 inches usable width
            try:
                det_table.columns[0].width = Inches(1.20)  # L1
                det_table.columns[1].width = Inches(1.22)  # V1
                det_table.columns[2].width = Inches(1.20)  # L2
                det_table.columns[3].width = Inches(1.22)  # V2
                det_table.columns[4].width = Inches(1.20)  # L3
                det_table.columns[5].width = Inches(1.22)  # V3
            except Exception:
                pass

            def _val(x):
                return '-' if (x is None or str(x).strip() == '') else str(x)

            # Helper to style a label cell
            def _style_label_cell(c):
                para = c.paragraphs[0]
                para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                run = para.runs[0]
                run.font.name = 'Calibri'
                run.font.size = Pt(8)
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
                set_cell_background(c, "44546A", RGBColor(255, 255, 255))

            # Helper to style a value cell
            def _style_value_cell(c):
                para = c.paragraphs[0]
                para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                if c.text:
                    run = para.runs[0]
                    run.font.name = 'Calibri'
                    run.font.size = Pt(8)

            # Populate rows
            for r, row_labels in enumerate(labels_grid):
                for pair_idx in range(3):
                    label = row_labels[pair_idx] if pair_idx < len(row_labels) else None
                    lcol = pair_idx * 2
                    if label:
                        # Label cell
                        lcell = det_table.cell(r, lcol)
                        lcell.text = f"{label}:"
                        _style_label_cell(lcell)
                        # Value cell
                        vcell = det_table.cell(r, lcol + 1)
                        vcell.text = _val(company_exposure.get(label))
                        _style_value_cell(vcell)
                    else:
                        # Empty pair - clear both cells
                        det_table.cell(r, lcol).text = ""
                        det_table.cell(r, lcol + 1).text = ""

            # Spacing after details table
            doc.add_paragraph()
    except Exception:
        # Fail silently if anything goes wrong with company details table
        pass

    # Build Key Credit Merits / Key Credit Risks table (to appear ABOVE the CAP table)
    # Build Key Credit Merits / Key Credit Risks table (to appear ABOVE the CAP table)
    try:
        merits = []
        risks = []

        def _find_key(d, names):
            """Find a value in dict d by any lowercase key name in names, allowing partial contains."""
            if not isinstance(d, dict):
                return None
            lower_map = {str(k).lower(): k for k in d.keys()}
            for n in names:
                if n in lower_map:
                    return d[lower_map[n]]
            for k in d.keys():
                kl = str(k).lower()
                for n in names:
                    if n in kl:
                        return d[k]
            return None

        def _to_list(x):
            """Convert various structures (dict with metric_1/risk_1, list, str) into a list of strings."""
            if x is None:
                return []
            if isinstance(x, list):
                return [str(i) for i in x if str(i).strip()]
            if isinstance(x, dict):
                def _idx(k):
                    m = re.search(r"(\d+)$", str(k))
                    return int(m.group(1)) if m else 999
                items = sorted(x.items(), key=lambda kv: _idx(kv[0]))
                return [str(v) for _, v in items if str(v).strip()]
            if isinstance(x, str):
                return [x] if x.strip() else []
            return [str(x)]

        src = credit_data
        if isinstance(src, dict) and src:
            root = _find_key(src, ["credit_risk_metrics"]) or src
            merits_raw = _find_key(root, ["key_credit_metrics", "key_credit_merits"])
            risks_raw = _find_key(root, ["key_credit_risks", "key_risks"])
            merits = _to_list(merits_raw)
            risks = _to_list(risks_raw)

        if merits or risks:
            max_rows = max(len(merits), len(risks))
            # Some spacing before the table
            doc.add_paragraph()
            # Create the 2-column table
            credit_table = doc.add_table(rows=max_rows + 1, cols=2)
            credit_table.style = 'Table Grid'

            # Header row styling
            hdr0 = credit_table.cell(0, 0)
            hdr1 = credit_table.cell(0, 1)
            headers = (('Key Credit Merits', hdr0), ('Key Credit Risks', hdr1))
            for text, cell in headers:
                cell.text = text
                para = cell.paragraphs[0]
                para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                run = para.runs[0]
                run.font.name = 'Calibri'
                run.font.size = Pt(8)
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
                set_cell_background(cell, "44546A", RGBColor(255, 255, 255))

            # Data rows
            for i in range(max_rows):
                ltxt = merits[i] if i < len(merits) else ""
                rtxt = risks[i] if i < len(risks) else ""
                lc = credit_table.cell(i + 1, 0)
                rc = credit_table.cell(i + 1, 1)
                lc.text = str(ltxt)
                rc.text = str(rtxt)
                for c in (lc, rc):
                    para = c.paragraphs[0]
                    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    if c.text:
                        run = para.runs[0]
                        run.font.name = 'Calibri'
                        run.font.size = Pt(8)

            # Column widths (split usable width roughly in half)
            credit_table.autofit = False
            credit_table.allow_autofit = False
            for i, column in enumerate(credit_table.columns):
                column.width = Inches(3.635)
            # Space after the table
            doc.add_paragraph()
    except Exception:
        # Fail silently if anything goes wrong with credit table
        pass
    
    # Helper function for formatting numbers in CAP table
    def _fmt_num(val):
        try:
            if val is None or val == "" or (isinstance(val, str) and val.strip() == "-"):
                return "-"
            # keep integers without decimals; floats with 2 decimals
            f = float(val)
            if abs(f - int(f)) < 1e-6:
                return f"{int(f):,}"
            return f"{f:,.2f}"
        except Exception:
            return str(val)
    
    # Add CAP table if available
    if isinstance(cap_json, dict):
        # First add a paragraph for spacing
        doc.add_paragraph()
        
        # Create CAP table with 6 columns
        cap_columns = ["Item", "Amount", "PPC Holdings", "Coupon", "Secured", "Maturity"]
        
        # Create a table with an extra row for the title
        cap_table = doc.add_table(rows=2, cols=len(cap_columns))
        cap_table.style = 'Table Grid'
        
        # Add title row that spans all columns
        title_row = cap_table.rows[0]
        title_cell = title_row.cells[0]
        # Merge all cells in the first row
        for i in range(1, len(cap_columns)):
            title_cell.merge(title_row.cells[i])
        
        # Set the title with company name
        title_cell.text = f"{company_title} - Capitalization Table"
        title_para = title_cell.paragraphs[0]
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_para.runs[0]
        title_run.font.name = 'Calibri'
        title_run.font.size = Pt(11)
        title_run.font.bold = True
        title_run.font.color.rgb = RGBColor(255, 255, 255)
        
        # Set background color for title
        title_shading = OxmlElement('w:shd')
        title_shading.set(qn('w:fill'), "44546A")  # Dark blue
        title_cell._element.tcPr.append(title_shading)
        
        # Add header row
        header_row = cap_table.rows[1]
        for i, col_name in enumerate(cap_columns):
            cell = header_row.cells[i]
            cell.text = col_name
            # Format header cell
            cell_para = cell.paragraphs[0]
            cell_para.alignment = WD_ALIGN_PARAGRAPH.CENTER if i > 0 else WD_ALIGN_PARAGRAPH.LEFT
            run = cell_para.runs[0]
            run.font.name = 'Calibri'
            run.font.size = Pt(8)
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
            # Set background color
            shading = OxmlElement('w:shd')
            shading.set(qn('w:fill'), "44546A")  # Dark blue
            cell._element.tcPr.append(shading)
        
        # (Removed explicit 'As of' row to match screenshot layout)
        
        # Cash and Equivalents
        cae = cap_json.get('cash_and_equivalents')
        row = cap_table.add_row()
        row.cells[0].text = "Cash and Equivalents"
        row.cells[1].text = _fmt_num(cae)
        
        # Debt breakdown
        debt_list = cap_json.get('debt') or []
        for d in debt_list:
            if not isinstance(d, dict):
                continue
            row = cap_table.add_row()
            row.cells[0].text = str(d.get('type', ''))
            row.cells[1].text = _fmt_num(d.get('amount'))
            row.cells[2].text = str(d.get('ppc_holdings', ''))
            row.cells[3].text = str(d.get('coupon', ''))
            row.cells[4].text = str(d.get('secured', ''))
            row.cells[5].text = str(d.get('maturity', ''))
        
        # Add important totals in a specific order
        for key, disp in [
            ("total_debt", "Total Debt"),
            ("book_value_of_equity", "Book Value of Equity"),
            ("book_capitalization", "Book Capitalization"),
            ("market_value_of_equity", "Market Value of Equity"),
            ("market_capitalization", "Market Capitalization"),
            ("ltm_adj_ebitda", "LTM Adjusted EBITDA"),
            ("market_value_of_re_assets", "Market Value of RE Assets"),
            ("unencumbered_assets", "Unencumbered Assets"),
        ]:
            val = cap_json.get(key)
            if val is not None:
                row = cap_table.add_row()
                row.cells[0].text = disp
                row.cells[1].text = _fmt_num(val)
                
                # Format special rows with bold text
                if disp in ["Total Debt", "Book Capitalization", "Market Capitalization"]:
                    for cell in row.cells[:2]:  # Only format the first two cells
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.font.bold = True
                    # Add a strong top border across the entire row as a separator
                    for cell in row.cells:
                        tc = cell._element.tcPr
                        if tc is None:
                            tc = OxmlElement('w:tcPr')
                            cell._element.append(tc)
                        borders = tc.find(qn('w:tcBorders'))
                        if borders is None:
                            borders = OxmlElement('w:tcBorders')
                            tc.append(borders)
                        top = OxmlElement('w:top')
                        top.set(qn('w:val'), 'single')
                        top.set(qn('w:sz'), '8')  # slightly thicker
                        top.set(qn('w:space'), '0')
                        top.set(qn('w:color'), '000000')
                        borders.append(top)
        
        # Key financial ratios header
        kfr = cap_json.get('key_financial_ratios') or {}
        if isinstance(kfr, dict) and kfr:
            row = cap_table.add_row()
            cell = row.cells[0]
            # Merge cells for the header
            for i in range(1, len(cap_columns)):
                cell.merge(row.cells[i])
            cell.text = "Key Financial Ratios:"
            # Format the cell
            cell_para = cell.paragraphs[0]
            cell_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = cell_para.runs[0]
            run.font.name = 'Calibri'
            run.font.size = Pt(8)
            run.font.bold = True
            run.italic = True
            # Set background color to light gray
            shading = OxmlElement('w:shd')
            shading.set(qn('w:fill'), "D3D3D3")  # Light gray
            cell._element.tcPr.append(shading)
            
            # Add ratio rows
            for k, v in kfr.items():
                label = k.replace('_', ' ').title() if isinstance(k, str) else str(k)
                row = cap_table.add_row()
                row.cells[0].text = label
                # Format numeric ratios; for select metrics, show percentage instead of 'x'
                def _norm(s: str) -> str:
                    return re.sub(r"\s+", " ", str(s).strip().lower())
                pct_labels = {
                    'total debt / book capital',
                    'total debt + leases / book capital',
                }
                is_pct_metric = _norm(label) in pct_labels
                display_v = v
                try:
                    if v is not None and str(v).strip() not in ('', '-'):
                        fv = float(str(v).replace('(', '-').replace(')', '').replace(',', '').replace('x', '').replace('%', ''))
                        if is_pct_metric:
                            display_v = f"({abs(fv):.1f}%)" if fv < 0 else f"{fv:.1f}%"
                        else:
                            display_v = f"({abs(fv):.2f}x)" if fv < 0 else f"{fv:.2f}x"
                except Exception:
                    display_v = v
                row.cells[1].text = str(display_v)
        
        # Format all cells in the table
        for row in cap_table.rows[1:]:  # Skip header row
            for i, cell in enumerate(row.cells):
                # Format cell text
                cell_para = cell.paragraphs[0]
                cell_para.alignment = WD_ALIGN_PARAGRAPH.CENTER if i > 0 else WD_ALIGN_PARAGRAPH.LEFT
                if cell.text:  # Only format if there's text
                    run = cell_para.runs[0]
                    run.font.name = 'Calibri'
                    run.font.size = Pt(8)  # Increased font size
        
        # Set column widths
        cap_table.autofit = False
        cap_table.allow_autofit = False
        for i, column in enumerate(cap_table.columns):
            if i == 0:
                column.width = Inches(2.5)  # First column wider
            else:
                column.width = Inches(4.77 / (len(cap_columns) - 1))  # Remaining width divided among other columns
        
        # Add space after CAP table
        doc.add_paragraph()
    
    # Add a page break before HFA table if CAP table exists
    if isinstance(cap_json, dict):
        doc.add_page_break()
    
    # Determine groups from columns for the two-row header
    columns = df.columns.tolist()
    year_cols = [c for c in columns if isinstance(c, str) and c.isdigit() and len(c) == 4]
    ytd_cols = [c for c in columns if isinstance(c, str) and c.startswith('YTD ')]
    ltm_cols = [c for c in columns if isinstance(c, str) and c.startswith('LTM ')]
    
    years_count = len(year_cols)
    ytd_count = len(ytd_cols)
    ltm_count = len(ltm_cols)
    
    # Create table with two header rows
    num_rows = len(df) + 2  # +2 for the two header rows
    num_cols = len(df.columns)
    
    # Now create the table
    table = doc.add_table(rows=num_rows, cols=num_cols)
    # Use a built-in style that's guaranteed to exist
    table.style = 'Table Grid'
    
    # Create two-row header with grouping labels
    # Row 1: First cell is empty, then "Fiscal Year Ended", "YTD", "LTM"
    header_row1 = table.rows[0]
    
    # First cell (Metric) - Company name and Historical Financial Analysis
    first_cell = header_row1.cells[0]
    first_cell.text = f"{company_title} - Historical Financial Analysis"
    first_para = first_cell.paragraphs[0]
    first_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    first_run = first_para.runs[0]
    first_run.font.name = 'Calibri'
    first_run.font.size = Pt(8)
    first_run.font.bold = True
    # Set background color to dark blue and text color to white
    set_cell_background(first_cell, "44546A", RGBColor(255, 255, 255))
    
    # Fiscal Year Ended group
    if years_count > 0:
        # Get the first and last year column index
        year_start_idx = 1  # First column after Metric
        year_end_idx = year_start_idx + years_count - 1
        
        # Merge cells for the Fiscal Year Ended header - correct way to merge cells
        fiscal_year_cell = header_row1.cells[year_start_idx]
        # We need to merge with each cell one by one in sequence
        for i in range(year_start_idx + 1, year_end_idx + 1):
            # Always merge with the next cell (which is now at position year_start_idx + 1 after each merge)
            fiscal_year_cell.merge(header_row1.cells[year_start_idx + 1])
        
        # Set text and formatting - ensure it's properly centered
        fiscal_year_cell.text = "Fiscal Year Ended"
        fiscal_year_para = fiscal_year_cell.paragraphs[0]
        fiscal_year_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Add tab stops to ensure proper centering
        tab_stops = fiscal_year_para.paragraph_format
        tab_stops.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fiscal_year_run = fiscal_year_para.runs[0]
        fiscal_year_run.font.name = 'Calibri'
        fiscal_year_run.font.size = Pt(8)
        fiscal_year_run.font.bold = True
        # Set background color to dark blue and text color to white
        set_cell_background(fiscal_year_cell, "44546A", RGBColor(255, 255, 255))
    
    # YTD group
    if ytd_count > 0:
        # Get the first and last YTD column index
        ytd_start_idx = year_start_idx + years_count
        ytd_end_idx = ytd_start_idx + ytd_count - 1
        
        # Merge cells for the YTD header - correct way to merge cells
        ytd_cell = header_row1.cells[ytd_start_idx]
        # We need to merge with each cell one by one in sequence
        for i in range(ytd_start_idx + 1, ytd_end_idx + 1):
            # Always merge with the next cell (which is now at position ytd_start_idx + 1 after each merge)
            ytd_cell.merge(header_row1.cells[ytd_start_idx + 1])
        
        # Set text and formatting
        ytd_cell.text = "YTD"
        ytd_para = ytd_cell.paragraphs[0]
        ytd_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        ytd_run = ytd_para.runs[0]
        ytd_run.font.name = 'Calibri'
        ytd_run.font.size = Pt(8)
        ytd_run.font.bold = True
        # Set background color to dark blue and text color to white
        set_cell_background(ytd_cell, "44546A", RGBColor(255, 255, 255))
    
    # LTM group
    if ltm_count > 0:
        # Get the first and last LTM column index
        ltm_start_idx = ytd_start_idx + ytd_count
        ltm_end_idx = ltm_start_idx + ltm_count - 1
        
        # Merge cells for the LTM header - correct way to merge cells
        ltm_cell = header_row1.cells[ltm_start_idx]
        # We need to merge with each cell one by one in sequence
        for i in range(ltm_start_idx + 1, ltm_end_idx + 1):
            # Always merge with the next cell (which is now at position ltm_start_idx + 1 after each merge)
            ltm_cell.merge(header_row1.cells[ltm_start_idx + 1])
        
        # Set text and formatting
        ltm_cell.text = "LTM"
        ltm_para = ltm_cell.paragraphs[0]
        ltm_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        ltm_run = ltm_para.runs[0]
        ltm_run.font.name = 'Calibri'
        ltm_run.font.size = Pt(8)
        ltm_run.font.bold = True
        # Set background color to dark blue and text color to white
        set_cell_background(ltm_cell, "44546A", RGBColor(255, 255, 255))
    
    # Row 2: Column headers
    header_row2 = table.rows[1]
    
    # First column is "Metric"
    metric_cell = header_row2.cells[0]
    metric_cell.text = "Metric"
    metric_para = metric_cell.paragraphs[0]
    metric_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    metric_run = metric_para.runs[0]
    metric_run.font.name = 'Calibri'
    metric_run.font.size = Pt(8)
    metric_run.font.bold = True
    # Set background color to dark blue and text color to white
    set_cell_background(metric_cell, "44546A", RGBColor(255, 255, 255))
    
    # Fill in the column headers
    col_idx = 1
    
    # Years
    for year in year_cols:
        cell = header_row2.cells[col_idx]
        cell.text = str(year)
        cell_para = cell.paragraphs[0]
        cell_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell_run = cell_para.runs[0]
        cell_run.font.name = 'Calibri'
        cell_run.font.size = Pt(8)
        cell_run.font.bold = True
        # Set background color to dark blue and text color to white
        set_cell_background(cell, "44546A", RGBColor(255, 255, 255))
        col_idx += 1
    
    # YTD dates
    for ytd in ytd_cols:
        cell = header_row2.cells[col_idx]
        try:
            yr = int(str(ytd).split()[1])
            cell.text = quarter_end_label_for_year(yr)
        except Exception:
            cell.text = str(ytd)
        cell_para = cell.paragraphs[0]
        cell_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell_run = cell_para.runs[0]
        cell_run.font.name = 'Calibri'
        cell_run.font.size = Pt(8)
        cell_run.font.bold = True
        # Set background color to dark blue and text color to white
        set_cell_background(cell, "44546A", RGBColor(255, 255, 255))
        col_idx += 1
    
    # LTM dates
    for ltm in ltm_cols:
        cell = header_row2.cells[col_idx]
        try:
            yr = int(str(ltm).split()[1])
            cell.text = quarter_end_label_for_year(yr)
        except Exception:
            cell.text = str(ltm)
        cell_para = cell.paragraphs[0]
        cell_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell_run = cell_para.runs[0]
        cell_run.font.name = 'Calibri'
        cell_run.font.size = Pt(8)
        cell_run.font.bold = True
        # Set background color to dark blue and text color to white
        set_cell_background(cell, "44546A", RGBColor(255, 255, 255))
        col_idx += 1
    
    # Track whether we've inserted the Key Financial Ratios header
    kfr_inserted = False

    # Fill in the data rows
    for i, row in enumerate(df.values):
        # Base index of the target row in the Word table (+2 header rows)
        # If we've already inserted a KFR header row, we need to shift all following rows by +1
        table_row_idx = i + 2 + (1 if kfr_inserted else 0)
        
        # Get the metric name (first column)
        metric_name = str(row[0]) if row[0] != '' else ''
        
        # Check if this row should have special formatting
        is_bold_row = metric_name in ['Revenue', 'Gross Profit', 'Adjusted EBITDA', 'Free Cash Flow', 
                                     'Total Debt', 'Book Equity', 'Change in Cash', 'Cash - End of Period']
        # Check for % Margin rows - these need special handling
        is_margin_row = metric_name == '% Margin' or metric_name.strip() == '%'
        is_growth_row = metric_name == '% YoY Growth'
        is_other_row = metric_name == 'Other'
        is_indent_row = is_margin_row or is_growth_row or is_other_row
        is_kfr_header = metric_name == 'EBITDA / Int. Exp.' or 'Key Financial Ratios:' in metric_name
        # Ratio row detection
        is_ratio_row = any(k in metric_name for k in ratio_keywords)
        pct_ratio_metrics = {
            'Total Debt / Book Capital',
            'Total Debt + Leases / Book Capital',
        }
        
        # Insert Key Financial Ratios header exactly once: just before the first ratio row encountered
        if not kfr_inserted and is_ratio_row:
            # Ensure there is space for the header row at current index
            table.add_row()
            # Prepare the header row at current computed index
            kfr_row = table.rows[table_row_idx]
            kfr_cell = kfr_row.cells[0]
            # Merge across entire width
            for col_idx in range(1, num_cols):
                try:
                    kfr_cell.merge(kfr_row.cells[1])
                except Exception:
                    break
            # Style the header cell
            kfr_cell.text = ""
            kfr_para = kfr_cell.paragraphs[0]
            kfr_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            kfr_para.paragraph_format.space_before = Pt(0)
            kfr_para.paragraph_format.space_after = Pt(0)
            kfr_run = kfr_para.add_run("Key Financial Ratios:")
            kfr_run.font.name = 'Calibri'
            kfr_run.font.size = Pt(8)
            kfr_run.bold = True
            kfr_run.italic = True
            kfr_run.font.color.rgb = RGBColor(0, 0, 0)
            kfr_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            kfr_row.height = Pt(14)
            # Background
            kfr_shading = OxmlElement('w:shd')
            kfr_shading.set(qn('w:fill'), "D3D3D3")
            kfr_cell._element.tcPr.append(kfr_shading)
            # Move write index to the next row (so current ratio row is placed under the header)
            table_row_idx += 1
            kfr_inserted = True
        
        # If this row is the 'Key Financial Ratios:' header, merge across the full width
        if metric_name.strip().startswith('Key Financial Ratios:'):
            kfr_row = table.rows[table_row_idx]
            kfr_cell = kfr_row.cells[0]
            
            # Merge across entire row
            for _ in range(1, num_cols):
                try:
                    kfr_cell.merge(kfr_row.cells[1])
                except Exception:
                    break
            # Style (ensure visible)
            kfr_cell.text = ""
            kfr_para = kfr_cell.paragraphs[0]
            kfr_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            kfr_para.paragraph_format.space_before = Pt(0)
            kfr_para.paragraph_format.space_after = Pt(0)
            kfr_run = kfr_para.add_run('Key Financial Ratios:')
            kfr_run.font.name = 'Calibri'
            kfr_run.font.size = Pt(8)
            kfr_run.bold = True
            kfr_run.italic = True
            kfr_run.font.color.rgb = RGBColor(0, 0, 0)
            kfr_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            kfr_row.height = Pt(14)
            # Light gray background
            kfr_shading = OxmlElement('w:shd')
            kfr_shading.set(qn('w:fill'), "D3D3D3")
            kfr_cell._element.tcPr.append(kfr_shading)
            # Proceed to next row
            continue

        # Fill in the data for this row
        for j, cell_text in enumerate(row):
            cell = table.cell(table_row_idx, j)
            text_out = str(cell_text) if cell_text != '' else ''
            # Apply ratio formatting for ratio rows, data columns only
            if j > 0 and is_ratio_row and text_out not in ['', '-']:
                try:
                    v = float(str(text_out).replace('(', '-').replace(')', '').replace(',', '').replace('x', '').replace('%', ''))
                    if metric_name in pct_ratio_metrics:
                        text_out = f"({abs(v):.1f}%)" if v < 0 else f"{v:.1f}%"
                    else:
                        text_out = f"({abs(v):.2f}x)" if v < 0 else f"{v:.2f}x"
                except Exception:
                    pass
            cell.text = text_out
            
            # Format data cell
            cell_para = cell.paragraphs[0]
            cell_para.alignment = WD_ALIGN_PARAGRAPH.CENTER if j > 0 else WD_ALIGN_PARAGRAPH.LEFT
            if cell.text:  # Only format if there's text
                run = cell_para.runs[0]
                run.font.name = 'Calibri'
                run.font.size = Pt(8)  # Increased font size
                
                # Apply bold formatting if needed
                if is_bold_row:
                    run.font.bold = True
                
                # Apply indentation if needed
                if j == 0:  # First column handling
                    if is_margin_row:
                        # Special handling for % Margin rows
                        cell.text = "% Margin"
                        # apply paragraph left indentation instead of leading spaces
                        cell_para = cell.paragraphs[0]
                        cell_para.paragraph_format.left_indent = Inches(0.15)
                        run = cell_para.runs[0]
                        run.font.name = 'Calibri'
                        run.font.size = Pt(8)
                        run.bold = True  # Make it bold
                        run.italic = True
                    elif is_growth_row:
                        # Special handling for % YoY Growth rows
                        cell.text = "% YoY Growth"
                        cell_para = cell.paragraphs[0]
                        cell_para.paragraph_format.left_indent = Inches(0.15)
                        run = cell_para.runs[0]
                        run.font.name = 'Calibri'
                        run.font.size = Pt(8)
                        run.bold = True  # Make it bold
                        run.italic = True
                    elif is_other_row:
                        # Special handling for Other rows
                        cell.text = "Other"
                        cell_para = cell.paragraphs[0]
                        cell_para.paragraph_format.left_indent = Inches(0.15)
                        run = cell_para.runs[0]
                        run.font.name = 'Calibri'
                        run.font.size = Pt(8)
                        run.italic = True
                    elif is_indent_row:
                        # General indentation
                        cell.text = f"{cell_text}"
                        cell_para = cell.paragraphs[0]
                        cell_para.paragraph_format.left_indent = Inches(0.15)
                        run = cell_para.runs[0]
                        run.font.name = 'Calibri'
                        run.font.size = Pt(8)
        
        # Add horizontal lines above specific rows
        if metric_name in ["Revenue", "Gross Profit", "Operating Expenses", "Adjusted EBITDA", 
                          "Interest Expense", "Capital Expenditures", "Free Cash Flow", 
                          "Acq. / Disp.", "Equity / Dividends", "Change in Cash", 
                          "Cash - End of Period", "Total Debt", "Book Equity"]:
            # Add a border to the top of this row
            for j in range(num_cols):
                cell = table.cell(table_row_idx, j)
                tc = cell._element.tcPr
                if tc is None:
                    tc = OxmlElement('w:tcPr')
                    cell._element.append(tc)
                
                # Add top border
                tcBorders = OxmlElement('w:tcBorders')
                tc.append(tcBorders)
                top = OxmlElement('w:top')
                top.set(qn('w:val'), 'single')
                top.set(qn('w:sz'), '4')
                top.set(qn('w:space'), '0')
                top.set(qn('w:color'), '000000')
                tcBorders.append(top)
                
                # Remove vertical borders
                for side in ['left', 'right']:
                    side_element = OxmlElement(f'w:{side}')
                    side_element.set(qn('w:val'), 'nil')
                    tcBorders.append(side_element)
    
    # Removed previous post-processing merge for 'Key Financial Ratios:' to avoid width overflow

    # Set column widths
    table.autofit = False
    table.allow_autofit = False
    for i, column in enumerate(table.columns):
        if i == 0:
            column.width = Inches(3.5)  # 45% of ~7.27 inches (A4 width minus margins)
        else:
            column.width = Inches(3.77 / (num_cols - 1))  # Remaining width divided among other columns
    
    # Set borders for the table - only vertical lines at left and right edges
    for row_idx, row in enumerate(table.rows):
        # detect KFR merged header row
        is_kfr_row = False
        try:
            first_txt = row.cells[0].text.strip()
            is_kfr_row = first_txt.startswith("Key Financial Ratios:")
        except Exception:
            is_kfr_row = False
        for cell_idx, cell in enumerate(row.cells):
            tc = cell._element.tcPr
            if tc is None:
                tc = OxmlElement('w:tcPr')
                cell._element.append(tc)
            
            # Ensure borders element exists
            borders = tc.find(qn('w:tcBorders'))
            if borders is None:
                borders = OxmlElement('w:tcBorders')
                tc.append(borders)
            
            # Set vertical borders only at the edges
            # Left border only for first column (for KFR row, only apply to first occurrence)
            left_element = OxmlElement('w:left')
            if cell_idx == 0:  # First column
                left_element.set(qn('w:val'), 'single')
                left_element.set(qn('w:sz'), '4')
                left_element.set(qn('w:space'), '0')
                left_element.set(qn('w:color'), '000000')
            elif is_kfr_row:
                # No inner verticals for merged KFR row
                left_element.set(qn('w:val'), 'nil')
            else:  # Other columns
                left_element.set(qn('w:val'), 'nil')
            borders.append(left_element)
            
            # Right border only for last column
            right_element = OxmlElement('w:right')
            if cell_idx == num_cols - 1:  # Last column
                right_element.set(qn('w:val'), 'single')
                right_element.set(qn('w:sz'), '4')
                right_element.set(qn('w:space'), '0')
                right_element.set(qn('w:color'), '000000')
            elif is_kfr_row:
                # No inner verticals for merged KFR row
                right_element.set(qn('w:val'), 'nil')
            else:  # Other columns
                right_element.set(qn('w:val'), 'nil')
            borders.append(right_element)
            
            # Ensure header cells have the proper background color
            if row_idx <= 1:  # First two rows are headers
                set_cell_background(cell, "44546A", RGBColor(255, 255, 255))
    
    # Add a page break after the table
    doc.add_paragraph().add_run().add_break()
    
    # Add Financial Statement Analysis from JSON (if present)
    if isinstance(fsa_data, dict):
        # doc.add_heading("Financial Statement Analysis", level=1)
        
        preferred_order = ["Income Statement", "Cash Flow Statement", "Balance Sheet"]
        for sec_name in preferred_order + [k for k in fsa_data.keys() if k not in preferred_order]:
            if sec_name in fsa_data and isinstance(fsa_data[sec_name], list) and fsa_data[sec_name]:
                # Add section header
                section_para = doc.add_heading(sec_name, level=2)
                section_para.runs[0].font.name = 'Calibri'
                section_para.runs[0].font.size = Pt(11)
                
                # Add bullet points
                for point in fsa_data[sec_name]:
                    bullet_para = doc.add_paragraph(point, style='List Bullet')
                    bullet_para.runs[0].font.name = 'Calibri'
                    bullet_para.runs[0].font.size = Pt(10)
    else:
        # doc.add_heading("Financial Statement Analysis", level=1)
        doc.add_paragraph("No statement analysis data found for this ticker.")
        
    # Add ESG Risk Ratings template table (empty data) after FSA
    try:
        # Define headers and ESG factor labels
        esg_headers = [
            "ESG Factor", "Risk Rating",
            "ESG Factor", "Risk Rating",
            "ESG Factor", "Risk Rating",
        ]
        left_esg = [
            "Climate Regulation",
            "Climate Change",
            "Habitat",
            "Sustainability",
            "Blended Score",
            "ESG Engagement",
        ]
        mid_esg = [
            "Product Safety",
            "Workplace Safety",
            "Health & Wellness",
            "Stakeholder Engagement",
            "Max Factor Score",
            "",
        ]
        right_esg = [
            "Board Composition",
            "Succession planning",
            "Data Security",
            "Labor Relations",
            "Aggregate Risk",
            "",
        ]
        max_rows = max(len(left_esg), len(mid_esg), len(right_esg))

        # Create the table: 1 header row + data rows, 6 columns
        esg_table = doc.add_table(rows=1 + max_rows, cols=6)
        esg_table.style = 'Table Grid'

        # Header row styling
        header_row = esg_table.rows[0]
        for j in range(6):
            hcell = header_row.cells[j]
            hcell.text = esg_headers[j]
            para = hcell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT if j % 2 == 0 else WD_ALIGN_PARAGRAPH.CENTER
            run = para.runs[0]
            run.font.name = 'Calibri'
            run.font.size = Pt(8)
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
            set_cell_background(hcell, "44546A", RGBColor(255, 255, 255))

        # Data rows
        for i in range(max_rows):
            row = esg_table.rows[1 + i]
            vals = [
                left_esg[i] if i < len(left_esg) else "",
                "",
                mid_esg[i] if i < len(mid_esg) else "",
                "",
                right_esg[i] if i < len(right_esg) else "",
                "",
            ]
            for j in range(6):
                cell = row.cells[j]
                # Put '*' in any empty ESG factor or rating cell
                cell.text = vals[j] if vals[j] else "*"
                para = cell.paragraphs[0]
                para.alignment = WD_ALIGN_PARAGRAPH.LEFT if j % 2 == 0 else WD_ALIGN_PARAGRAPH.CENTER
                if cell.text:
                    run = para.runs[0]
                    run.font.name = 'Calibri'
                    run.font.size = Pt(8)
                    # Bold ESG Factor labels
                    if j % 2 == 0 and vals[j]:
                        run.bold = True

        # Optional: set column widths (approximate)
        try:
            esg_table.autofit = False
            esg_table.allow_autofit = False
            for idx, col in enumerate(esg_table.columns):
                if idx % 2 == 0:
                    col.width = Inches(1.75)
                else:
                    col.width = Inches(0.65)
        except Exception:
            pass

        # Spacing after table
        doc.add_paragraph()
    except Exception:
        # Do not fail document generation if ESG table build fails
        pass

    # Add title at the end of the document
    doc.add_paragraph()
    title_para = doc.add_paragraph()
    # title_run = title_para.add_run(f"{company_title} - Historical Financial Analysis")
    title_run.font.name = 'Calibri'
    title_run.font.size = Pt(14)
    title_run.font.bold = True
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add COMP table if available
    if isinstance(comp_rows, list) and comp_rows:
        try:
            # Add a page break before the Comparables Analysis section
            doc.add_page_break()
            
            # Convert COMP rows to DataFrame
            df_comp = json_to_dataframe(comp_rows)
            df_comp = df_comp.replace({np.nan: '-'})  # Replace NaN with dash for better display
            
            # Define ticker to company name mapping
            ticker_to_company = {
                'ELME': 'Elme Communities',
                'MAA': 'Mid-America Apartment Communities',
                'CPT': 'Camden Property Trust',
                'UDR': 'UDR, Inc.',
                'IRT': 'Independence Realty Trust',
                'EQR': 'Equity Residential',
                'ESS': 'Essex Property Trust',
                'AVB': 'AvalonBay Communities',
                'AIV': 'Apartment Investment and Management',
                'AVERAGE': 'Average',
                'MEDIAN': 'Median'
            }
            
            # Format revenue and EBITDA columns by removing trailing zeros (000s)
            # First identify which columns are LTM Rev and LTM EBITDA
            rev_col_idx = None
            ebitda_col_idx = None
            
            # Check column names
            for i, col in enumerate(df_comp.columns):
                col_str = str(col).upper()
                # Exact match for LTM EBITDA
                if col_str == "LTM EBITDA":
                    ebitda_col_idx = i
                # Look for revenue column
                elif any(term in col_str for term in ['REV', 'REVENUE']):
                    rev_col_idx = i
                # Look for EBITDA column - check both spellings
                elif any(term in col_str for term in ['EBITDA', 'EBTDA']) and not any(term in col_str for term in ['MARGIN', 'MRGN', '%']):
                    ebitda_col_idx = i
            
            # If not found, use default indices
            if rev_col_idx is None:
                rev_col_idx = 1  # Second column (index 1) is typically LTM Rev
            if ebitda_col_idx is None:
                ebitda_col_idx = 2  # Third column (index 2) is typically LTM EBITDA

            # Build adjusted headers based on df_comp (remove Rating/As of)
            # Mapping from df_comp column names to display headers
            column_mapping = {
                "LTM Revenue": "Revenue",
                "LTM EBITDA": "LTM EBITDA",
                "EBITDA Margin %": "EBITDA Margin %",
                "EBITDAR / (Int + Rents)": "EBITDAR / (Int + Rents)",
                "(Total Debt + COL) / EBITDAR": "Total Debt + COL / EBITDAR",
                "(Net Debt + COL) / EBITDAR": "Net Debt + COL / EBITDAR",
                "(Total Debt + COL) / Total Cap": "Total Debt + COL / Total Cap",
                "(FCF + Rents) / (Total Debt + COL)": "FCF + Rents / Total Debt + COL",
                "3Y Avg (TD+COL)/EBITDAR": "3-Year Average Total Debt + COL / EBITDAR",
                "3Y Avg (TD+COL)/Total Cap": "3-Year Average Total Debt + COL / Total Cap",
                "3Y Avg (FCF+Rents)/(TD+COL)": "3-Year Average FCF + Rents / Total Debt + COL"
            }

            adjusted_headers = ["Ticker"]
            # Map df_comp index to our header name
            df_to_header_map: dict[int, str] = {}
            for i, col in enumerate(df_comp.columns):
                col_str = str(col)
                if col_str in column_mapping:
                    header_name = column_mapping[col_str]
                    df_to_header_map[i] = header_name
                    if header_name not in adjusted_headers:
                        adjusted_headers.append(header_name)

            # Create the comp table sized exactly to the adjusted headers
            # We add an extra group-header row (LTM / 3-Year Average)
            num_rows = len(df_comp) + 3  # +3 for title row, group row, and column header row
            num_cols = len(adjusted_headers)
            comp_table = doc.add_table(rows=num_rows, cols=num_cols)
            comp_table.style = 'Table Grid'

            # Add title row that spans all columns
            title_row = comp_table.rows[0]
            title_cell = title_row.cells[0]
            # Merge all cells in the first row
            for i in range(1, num_cols):
                title_cell.merge(title_row.cells[i])

            # Set the title with company name
            title_cell.text = f"{company_title} - Credit Comparable Analysis"
            title_para = title_cell.paragraphs[0]
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_run = title_para.runs[0]
            title_run.font.name = 'Calibri'
            title_run.font.size = Pt(11)
            title_run.font.bold = True

            # Set background color for title to light gray
            title_shading = OxmlElement('w:shd')
            title_shading.set(qn('w:fill'), "D3D3D3")  # Light gray
            title_cell._element.tcPr.append(title_shading)

            # Note: adjusted_headers and df_to_header_map already built above

            # Build a dark group header row for LTM and 3-Year Average (row index 1)
            group_row = comp_table.rows[1]
            # helper to dark-style a header cell
            def _dark_header(cell):
                para = cell.paragraphs[0]
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = para.runs[0] if para.runs else para.add_run("")
                run.font.name = 'Calibri'
                run.font.size = Pt(9)
                run.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
                set_cell_background(cell, "44546A", RGBColor(255, 255, 255))
            # initialize group cells
            for i in range(num_cols):
                group_row.cells[i].text = ""
                _dark_header(group_row.cells[i])
            # compute contiguous spans for LTM group and 3-Year Average group
            def _idx(name):
                return adjusted_headers.index(name) if name in adjusted_headers else None
            ltm_members = [
                "Revenue", "LTM EBITDA", "EBITDA Margin %", "EBITDAR / (Int + Rents)",
                "Total Debt + COL / EBITDAR", "Net Debt + COL / EBITDAR",
                "Total Debt + COL / Total Cap", "FCF + Rents / Total Debt + COL"
            ]
            ltm_indices = sorted([_idx(h) for h in ltm_members if _idx(h) is not None])
            if ltm_indices:
                first, last = ltm_indices[0], ltm_indices[-1]
                merged = group_row.cells[first]
                for _ in range(first + 1, last + 1):
                    merged.merge(group_row.cells[first + 1])
                merged.text = "LTM"
                _dark_header(merged)
            avg_members = [
                "3Y Avg (TD+COL)/EBITDAR",
                "3Y Avg (TD+COL)/Total Cap",
                "3Y Avg (FCF+Rents)/(TD+COL)"
            ]
            avg_indices = sorted([_idx(h) for h in avg_members if _idx(h) is not None])
            if avg_indices:
                first, last = avg_indices[0], avg_indices[-1]
                merged = group_row.cells[first]
                for _ in range(first + 1, last + 1):
                    merged.merge(group_row.cells[first + 1])
                merged.text = "3-Year Average"
                _dark_header(merged)

            # Column header row (row index 2)
            header_row = comp_table.rows[2]
            for i, header in enumerate(adjusted_headers):
                cell = header_row.cells[i]
                # display labels with line breaks (keep spelling 'EBITDA')
                display = header
                if header == "Revenue":
                    display = "LTM Rev\n(000s)"
                elif header == "LTM EBITDA":
                    display = "LTM EBITDA\n(000s)"
                elif header == "EBITDA Margin %":
                    display = "EBITDA\nMrgn %"
                elif header == "EBITDAR / (Int + Rents)":
                    display = "EBITDAR\n(Int + Rents)"
    
                elif header == "Total Debt + COL / EBITDAR":
                    display = "(Tot Debt + COL)\nEBITDAR"
                elif header == "Net Debt + COL / EBITDAR":
                    display = "(Net Debt + COL)\nEBITDAR"
                elif header == "Total Debt + COL / Total Cap":
                    display = "(Tot Debt + COL)\nTot Cap"
                elif header == "FCF + Rents / Total Debt + COL":
                    display = "(FCF + Rents)\n(Tot Debt + COL)"
                elif header == "3Y Avg (TD+COL)/EBITDAR":
                    display = "3Y Avg\n(TD+COL) EBITDAR"
                elif header == "3Y Avg (TD+COL)/Total Cap":
                    display = "3Y Avg\n(TD+COL) Tot Cap"
                elif header == "3Y Avg (FCF+Rents)/(TD+COL)":
                    display = "3Y Avg\n(FCF+Rents) (TD+COL)"
                elif i == 0:
                    display = "Ticker"

                para = cell.paragraphs[0]
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER if i > 0 else WD_ALIGN_PARAGRAPH.LEFT
                # clear any default run text
                if para.runs:
                    para.runs[0].text = display
                    run = para.runs[0]
                else:
                    run = para.add_run(display)
                run.font.name = 'Calibri'
                run.font.size = Pt(8)
                run.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
                set_cell_background(cell, "44546A", RGBColor(255, 255, 255))

            # Fill in the data rows
            for i, row in enumerate(df_comp.values):
                table_row = comp_table.rows[i + 3]  # +3 to account for title + group + header rows

                # First column: Ticker
                ticker = str(row[0]).strip().upper() if row[0] != '' else ''
                table_row.cells[0].text = ticker

                # Format company name cell
                cell_para = table_row.cells[0].paragraphs[0]
                cell_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                run = cell_para.runs[0]
                run.font.name = 'Calibri'
                run.font.size = Pt(8)
                if ticker.upper() not in ("AVERAGE", "MEDIAN"):
                    run.italic = True
                else:
                    run.font.bold = True
                    # Add light gray background for Average and Median rows
                    shading = OxmlElement('w:shd')
                    shading.set(qn('w:fill'), "D3D3D3")  # Light gray
                    table_row.cells[0]._element.tcPr.append(shading)

                # Fill in the data for the remaining columns
                for j, cell_text in enumerate(row):
                    if j == 0:  # Skip ticker column as we already handled it
                        continue

                    # Map df_comp column index to our header name
                    if j not in df_to_header_map:
                        continue
                    header_name = df_to_header_map[j]
                    # Find the position in our table
                    table_col_idx = adjusted_headers.index(header_name)
                    if table_col_idx < 0 or table_col_idx >= num_cols:
                        continue

                    # Format the cell value
                    cell_text = str(cell_text) if cell_text != '' else '-'

                    # Format numbers based on column type
                    if cell_text != '-':
                        try:
                            val = float(str(cell_text).replace(',', ''))
                            header = adjusted_headers[table_col_idx]
                            
                            # Revenue and EBITDA: divide by 1000 and format with 1 decimal
                            if header == "Revenue" or header == "LTM EBITDA":
                                # val = val / 1000
                                cell_text = f"{val:.1f}"
                            # Percentages: format with 1 decimal and % symbol
                            elif "Margin" in header or "%" in header:
                                cell_text = f"{val:.1f}%"
                            # Ratios: format with 2 decimals and x suffix
                            elif "EBITDAR / (Int + Rents)" in header or "FCF + Rents / Total Debt + COL" in header or "/" in header or "x" in header or "Ratio" in header:
                                cell_text = f"{val:.2f}x"
                            # Other metrics: keep 1 decimal (leave units as-is); if these are ratios without explicit match, also show x with 2 decimals
                            else:
                                cell_text = f"{val:.2f}x"
                        except (ValueError, TypeError):
                            pass  # Keep as is if not a number
                    
                    # Set the cell text
                    table_row.cells[table_col_idx].text = cell_text
                    
                    # Format the cell
                    cell_para = table_row.cells[table_col_idx].paragraphs[0]
                    cell_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = cell_para.runs[0]
                    run.font.name = 'Calibri'
                    run.font.size = Pt(8)  # Increased font size
                    
                    # Bold for Average and Median rows
                    if ticker.upper() in ("AVERAGE", "MEDIAN"):
                        run.font.bold = True
                        # Add light gray background
                        shading = OxmlElement('w:shd')
                        shading.set(qn('w:fill'), "D3D3D3")  # Light gray
                        table_row.cells[table_col_idx]._element.tcPr.append(shading)
            
            # Set column widths (Ticker ~15%, remaining evenly) and center the table
            comp_table.autofit = False
            comp_table.allow_autofit = False
            comp_table.alignment = WD_TABLE_ALIGNMENT.CENTER
            total_w_in = (section.page_width.inches - section.left_margin.inches - section.right_margin.inches) * 0.995
            first_col_frac = 0.15
            remaining_cols = max(1, num_cols - 1)
            other_frac = (1.0 - first_col_frac) / remaining_cols
            for i, column in enumerate(comp_table.columns):
                frac = first_col_frac if i == 0 else other_frac
                width_in = total_w_in * frac
                set_column_preferred_width(comp_table, i, width_in)
            # Force total table width
            set_table_fixed_width(comp_table, total_w_in)
            set_table_indent(comp_table, 0.0)
        except Exception as e:
            # If there's an error creating the COMP table, add a note
            doc.add_heading("Comparables Analysis", level=1)
            doc.add_paragraph(f"Error creating comparables table: {e}")
    
    # Covenant Summary Table (empty template) after COMP
    try:
        cov_company_title = get_company_title_from_sec(ticker) or company_title
        cov_title = f"{cov_company_title} - Covenant Summary"
        cov_date = "3/31/2025"

        # Build table: 3 columns (Term | Covenant Level | Reported)
        terms = [
            "Maximum Leverage Ratio",
            "Unconsolidated Affiliates / Total Asset Value",
            "Total Marketable Securities, etc. / Total Asset Value",
            "Minimum Fixed Charge Coverage Ratio",
            "Maximum Secured Indebtedness",
            "Maximum Unencumbered Leverage Ratio",
        ]
        more_terms = [
            "Unimprovement Land / Unencumbered Pool Value",
            "Development, JVs, etc. / Unencumbered Pool Value",
        ]

        rows = 2 + 1 + len(terms) + 1 + len(more_terms)  # title + date + header + terms + group + additional
        cov_table = doc.add_table(rows=rows, cols=3)
        cov_table.style = 'Table Grid'

        r = 0
        # Title row spanning 3 cols
        title_cell = cov_table.cell(r, 0)
        for c in range(1, 3):
            title_cell.merge(cov_table.cell(r, c))
        title_cell.text = cov_title
        para = title_cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.runs[0]
        run.font.name = 'Calibri'
        run.font.size = Pt(9)
        run.font.bold = True
        set_cell_background(title_cell, "44546A", RGBColor(255, 255, 255))
        r += 1

        # Date row spanning 3 cols
        date_cell = cov_table.cell(r, 0)
        for c in range(1, 3):
            date_cell.merge(cov_table.cell(r, c))
        date_cell.text = cov_date
        para = date_cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.runs[0]
        run.font.name = 'Calibri'
        run.font.size = Pt(9)
        run.bold = True
        # light gray background
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), "D3D3D3")
        date_cell._element.tcPr.append(shading)
        r += 1

        # Header row
        headers = ["Term", "Covenant Level", "Reported"]
        for c in range(3):
            hcell = cov_table.cell(r, c)
            hcell.text = headers[c]
            para = hcell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.runs[0]
            run.font.name = 'Calibri'
            run.font.size = Pt(8)
            run.bold = True
            # light gray background
            shading = OxmlElement('w:shd')
            shading.set(qn('w:fill'), "D3D3D3")
            hcell._element.tcPr.append(shading)
        r += 1

        # Term rows
        for term in terms:
            cells = [cov_table.cell(r, 0), cov_table.cell(r, 1), cov_table.cell(r, 2)]
            cells[0].text = term
            # Fill empty Covenant Level and Reported with '*'
            cells[1].text = "*"
            cells[2].text = "*"
            for idx, cell in enumerate(cells):
                para = cell.paragraphs[0]
                para.alignment = WD_ALIGN_PARAGRAPH.LEFT if idx == 0 else WD_ALIGN_PARAGRAPH.CENTER
                if cell.text:
                    run = para.runs[0]
                    run.font.name = 'Calibri'
                    run.font.size = Pt(8)
                    if idx == 0:
                        run.bold = True
            r += 1

        # Group header row spanning 3 cols
        group_cell = cov_table.cell(r, 0)
        for c in range(1, 3):
            group_cell.merge(cov_table.cell(r, c))
        group_cell.text = "Additional Covenants / Baskets"
        para = group_cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.runs[0]
        run.font.name = 'Calibri'
        run.font.size = Pt(8)
        run.bold = True
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), "D3D3D3")
        group_cell._element.tcPr.append(shading)
        r += 1

        for term in more_terms:
            cells = [cov_table.cell(r, 0), cov_table.cell(r, 1), cov_table.cell(r, 2)]
            cells[0].text = term
            # Fill empty Covenant Level and Reported with '*'
            cells[1].text = "*"
            cells[2].text = "*"
            for idx, cell in enumerate(cells):
                para = cell.paragraphs[0]
                para.alignment = WD_ALIGN_PARAGRAPH.LEFT if idx == 0 else WD_ALIGN_PARAGRAPH.CENTER
                if cell.text:
                    run = para.runs[0]
                    run.font.name = 'Calibri'
                    run.font.size = Pt(8)
                    if idx == 0:
                        run.bold = True
            r += 1

        # Force full-width table and proper alignment
        try:
            cov_table.autofit = False
            cov_table.allow_autofit = False
            cov_table.alignment = WD_TABLE_ALIGNMENT.CENTER
            total_w_in = 7.15  # slightly under full width to avoid Word auto-shrink
            widths_in = [total_w_in * 0.62, total_w_in * 0.19, total_w_in * 0.19]
            for idx, w_in in enumerate(widths_in):
                set_column_preferred_width(cov_table, idx, w_in)
            # Force total table width
            set_table_fixed_width(cov_table, total_w_in)
            set_table_indent(cov_table, 0.0)
        except Exception:
            pass
        # spacing
        doc.add_paragraph()
        # Add bottom note explaining '*'
        note_para = doc.add_paragraph("Note: '*' indicates the data source are private.")
        if note_para.runs:
            nrun = note_para.runs[0]
        else:
            nrun = note_para.add_run("")
            nrun.text = "Note: '*' indicates the data source are private."
        nrun.font.name = 'Calibri'
        nrun.font.size = Pt(8)
        nrun.font.color.rgb = RGBColor(128, 128, 128)
    except Exception:
        pass

    # Add a final page break if needed
    if isinstance(comp_rows, list) and comp_rows:
        # We already added a page break before the COMP table
        pass
    else:
        # Add a page break before the end of the document
        doc.add_page_break()
    
    # Save the document to the buffer
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


if __name__ == '__main__':
    parser = argparse.ArgumentParser(description='Generate a financial Word document for a ticker using HFA and FSA JSON files and save it locally.')
    parser.add_argument('-t', '--ticker', help='Ticker symbol used to locate JSON files (e.g., ELME)')
    parser.add_argument('-o', '--output', help='Output Word document filename (default: <ticker>_AQRR_{year}.docx)')
    args = parser.parse_args()

    ticker = args.ticker.strip() if args.ticker else input('Enter ticker symbol: ').strip()
    if not ticker:
        print('Error: Ticker is required.')
        sys.exit(1)

    try:
        word_bytes = build_word_bytes_from_ticker(ticker)
    except (FileNotFoundError, ValueError, RuntimeError) as e:
        print(f'Error: {e}')
        sys.exit(1)

    # Determine output path and filename
    year = datetime.now().year
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    default_dir = os.path.join('output', 'word', 'AQRR')
    os.makedirs(default_dir, exist_ok=True)
    
    # Create a unique filename with timestamp
    if args.output:
        out_path = args.output
    else:
        default_filename = f"{ticker}_AQRR_{year}_{timestamp}.docx"
        out_path = os.path.join(default_dir, default_filename)
    
    # Try to save the file
    max_attempts = 3
    attempt = 0
    success = False
    
    while attempt < max_attempts and not success:
        try:
            with open(out_path, 'wb') as f:
                f.write(word_bytes)
            print(f'Saved Word document to {os.path.abspath(out_path)}')
            success = True
        except PermissionError:
            attempt += 1
            if attempt < max_attempts:
                print(f"Permission denied. File may be open in another application. Trying alternative filename...")
                # Generate a new filename with a different timestamp
                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S") + f"_{attempt}"
                default_filename = f"{ticker}_AQRR_{year}_{timestamp}.docx"
                out_path = os.path.join(default_dir, default_filename)
            else:
                print(f"Error: Permission denied after {max_attempts} attempts. Please close any open Word documents and try again.")
                print(f"Alternatively, specify a different output path using the -o option.")
                sys.exit(1)
        except Exception as e:
            print(f'Failed to write output file: {e}')
            print(f"Please check that you have write permissions to {os.path.dirname(os.path.abspath(out_path))}")
            sys.exit(1)
